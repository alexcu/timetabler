import os

import pandas
from docx import Document
from pandas import ExcelFile
from pulp import LpProblem, LpMinimize, lpSum, LpVariable, LpStatus, LpInteger, LpBinary

from Attendance import app, db, executor
from Attendance.models import *


#TIMETABLE CODE
def runtimetable(STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
                 TUTORAVAILABILITY, maxclasssize, minclasssize, nrooms):
    '''
    Run the timetabling process and input into the database.

    This process calls the CBCSolver using the PuLP package and then adds the classes to the database.


    :param STUDENTS: should be an array of student names
    :param SUBJECTS: should be an array of subject codes
    :param TIMES: an array of strings representing possible timeslots
    :param day:
    :param DAYS: the days corresponding to the timeslots above
    :param TEACHERS: an array of the names of the tutors
    :param SUBJECTMAPPING: This is a dictionary representing the subjects
                            each tutor is taking
    :param REPEATS: A dictionary of how many repeats each subject has
    :param TEACHERMAPPING: A dictionary of what subject each tutor teachers
    :param TUTORAVAILABILITY:
    :param maxclasssize: An integer representing the maximum class size
    :param minclasssize: An integer representing the minimum class size
    :param nrooms: An integer representing the max allowable concurrent classes
    :return: A string representing model status.
    '''
    print("Running solver")
    model = LpProblem('Timetabling', LpMinimize)
    #Create Variables
    print("Creating Variables")
    assign_vars = LpVariable.dicts("StudentVariables",[(i, j,k,m) for i in STUDENTS for j in SUBJECTS for k in TIMES for m in TEACHERS], 0, 1, LpBinary)
    subject_vars = LpVariable.dicts("SubjectVariables",[(j,k,m) for j in SUBJECTS for k in TIMES for m in TEACHERS], 0, 1, LpBinary)

    #c
    num930classes = LpVariable.dicts("930Classes", [(i) for i in TIMES], lowBound = 0, cat = LpInteger)
    #w
    daysforteachers = LpVariable.dicts("numdaysforteachers", [(i,j) for i in TEACHERS for j in range(len(DAYS))], 0,1,LpBinary)
    #p
    daysforteacherssum = LpVariable.dicts("numdaysforteacherssum", [(i) for i in TEACHERS],0,cat = LpInteger)
    #variables for student clashes
    studenttime = LpVariable.dicts("StudentTime", [(i,j) for i in STUDENTS for j in TIMES],lowBound=0,upBound=1,cat=LpBinary)
    studentsum = LpVariable.dicts("StudentSum", [(i) for i in STUDENTS],0,cat = LpInteger)

    #Count the days that a teacher is rostered on. Make it bigger than a small number times the sum
    #for that particular day.
    for m in TEACHERS:
        for d in range(len(day)):
            model += daysforteachers[(m, d)] >= 0.1 * lpSum(
                subject_vars[(j, k, m)] for j in SUBJECTS for k in DAYS[day[d]])
            model += daysforteachers[(m, d)] <= lpSum(subject_vars[(j, k, m)] for j in SUBJECTS for k in DAYS[day[d]])
    for m in TEACHERS:
        model += daysforteacherssum[(m)] == lpSum(daysforteachers[(m, d)] for d in range(len(day)))

    print("Constraining tutor availability")
    #This bit of code puts in the constraints for the tutor availability.
    #It reads in the 0-1 matrix of tutor availability and constrains that no classes
    # can be scheduled when a tutor is not available.
    #The last column of the availabilities is the tutor identifying number, hence why we have
    #used a somewhat convoluted idea down here.
    for m in TEACHERS:
        for k in TIMES:
            if k not in TUTORAVAILABILITY[m]:
                model += lpSum(subject_vars[(j,k,m)] for j in SUBJECTS) == 0


    #Constraints on subjects for each students
    print("Constraining student subjects")
    for i in STUDENTS:
        for j in SUBJECTS:
            if i in SUBJECTMAPPING[j]:
                model += lpSum(assign_vars[(i,j,k,m)] for k in TIMES for m in TEACHERS) == 1
            else:
                model += lpSum(assign_vars[(i,j,k,m)] for k in TIMES for m in TEACHERS) == 0



    #This code means that students cannot attend a tute when a tute is not running
    #But can not attend a tute if they attend a repeat.
    for i in STUDENTS:
        for j in SUBJECTS:
            for k in TIMES:
                for m in TEACHERS:
                    model += assign_vars[(i,j,k,m)] <= subject_vars[(j,k,m)]


    #Constraints on which tutor can take each class
    #This goes through each list and either constrains it to 1 or 0 depending if
    #the teacher needs to teach that particular class.
    print("Constraining tutor classes")
    for m in TEACHERS:
        for j in SUBJECTS:
            if j in TEACHERMAPPING[m]:
                #THIS WILL BE CHANGED TO NUMBER OF REPEATS
                model+= lpSum(subject_vars[(j,k,m)] for k in TIMES) == REPEATS[j]
            else:
                model += lpSum(subject_vars[(j,k,m)] for k in TIMES) == 0

    #General Constraints on Rooms etc.
    print("Constraining times")
    # For each time cannot exceed number of rooms
    for k in TIMES:
        model += lpSum(subject_vars[(j,k,m)] for j in SUBJECTS for m in TEACHERS) <= nrooms

    #Teachers can only teach one class at a time
    for k in TIMES:
        for m in TEACHERS:
            model += lpSum(subject_vars[(j,k,m)] for j in SUBJECTS) <= 1

    #STUDENT CLASHES
    for i in STUDENTS:
        for k in TIMES:
            model += studenttime[(i,k)] <= lpSum(assign_vars[(i,j,k,m)] for j in SUBJECTS for m in TEACHERS)/2
            model += studenttime[(i, k)] >= 0.3*(0.5*lpSum(assign_vars[(i, j, k, m)] for j in SUBJECTS for m in TEACHERS) -0.5)

    for i in STUDENTS:
        model += studentsum[(i)] == lpSum(studenttime[(i,k)] for k in TIMES)

    #This minimizes the number of 9:30 classes.
    for i in TIMES:
        if i.find('21:30') != -1:
            model += num930classes[(i)] == lpSum(subject_vars[(j,i,m)] for j in SUBJECTS for m in TEACHERS)

        else:
            model += num930classes[(i)] == 0

    print("Setting objective function")

    #Class size constraint
    for j in SUBJECTS:
        for k in TIMES:
            for m in TEACHERS:
                model +=lpSum(assign_vars[(i,j,k,m)] for i in STUDENTS) >= minclasssize*subject_vars[(j,k,m)]
                model += lpSum(assign_vars[(i,j,k,m)] for i in STUDENTS) <= maxclasssize

    #Solving the model
    model += (100*lpSum(studentsum[(i)] for i in STUDENTS) + lpSum(num930classes[(i)] for i in TIMES) + 500*lpSum(daysforteacherssum[(m)] for m in TEACHERS))
    print("Solving Model")
    model.solve()
    print("Status:", LpStatus[model.status])
    print("Complete")
    for j in SUBJECTS:
        subject = Subject.get(subcode=j)
        for k in TIMES:
            timesplit = k.split(' ')
            timeslot = Timeslot.get(timetable=get_current_timetable().id, day=timesplit[0], time=timesplit[1])
            for m in TEACHERS:
                tutor = Tutor.get(name=m)
                if subject_vars[(j, k, m)].varValue == 1:
                    timetabledclass = TimetabledClass.create(subjectid=subject.id, timetable=get_current_timetable().id,
                                                             time=timeslot.id, tutorid=tutor.id)
                    for i in STUDENTS:
                        if assign_vars[(i, j, k, m)].varValue == 1:
                            student = Student.get(name=i)
                            timetabledclass.students.append(student)
                            db.session.commit()
    print("Status:", LpStatus[model.status])
    return LpStatus[model.status]


def preparetimetable(addtonewtimetable=False):
    '''
    Get timetable data and then execute the timetabling program.

    :param addtonewtimetable: Whether this should be added to a new timetable and set as default.
    :return: The view timetable page.
    '''
    # if addtonewtimetable == "true":
    #    timetable = Timetable(get_current_year(),get_current_studyperiod())
    #    db.session.add(timetable)
    #    db.session.commit()
    #    admin = Admin.query.filter_by(key="timetable").first()
    #    admin.value = timetable.id
    #    db.session.commit()
    print("Preparing Timetable")
    (STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
     TUTORAVAILABILITY, maxclasssize, minclasssize, nrooms) = get_timetable_data()
    print("Everything ready")
    executor.submit(runtimetable,STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
                       TUTORAVAILABILITY, maxclasssize, minclasssize, nrooms)
    return render_template("viewtimetable.html")


def get_timetable_data():
    '''
    Get all required timetable data from the database
    :return: All timetabling data as a tuple to the preparetimetable method.
    '''
    SUBJECTS = []
    SUBJECTMAPPING = {}
    STUDENTS = []
    REPEATS = {}
    TEACHERS = []
    TUTORAVAILABILITY = {}
    TEACHERMAPPING = {}
    allsubjects = Subject.query.filter(Subject.year == get_current_year(),
                                       Subject.studyperiod == get_current_studyperiod(), Subject.tutor != None).all()
    alltutors = []
    for subject in allsubjects:
        SUBJECTMAPPING[subject.subcode] = []
        REPEATS[subject.subcode] = subject.repeats
        SUBJECTS.append(subject.subcode)
        TEACHERS.append(subject.tutor.name)
        if subject.tutor not in alltutors:
            alltutors.append(subject.tutor)
        for student in subject.students:
            STUDENTS.append(student.name)
            SUBJECTMAPPING[subject.subcode].append(student.name)
        SUBJECTMAPPING[subject.subcode] = set(SUBJECTMAPPING[subject.subcode])
    STUDENTS = list(set(STUDENTS))
    TEACHERS = list(set(TEACHERS))
    for tutor in alltutors:
        TUTORAVAILABILITY[tutor.name] = []
        TEACHERMAPPING[tutor.name] = []
        for timeslot in tutor.availabletimes:
            TUTORAVAILABILITY[tutor.name].append(timeslot.day + " " + timeslot.time)
        for subject in tutor.subjects:
            TEACHERMAPPING[tutor.name].append(subject.subcode)
        TUTORAVAILABILITY[tutor.name] = set(TUTORAVAILABILITY[tutor.name])
        TEACHERMAPPING[tutor.name] = set(TEACHERMAPPING[tutor.name])

    maxclasssize = 400
    minclasssize = 1
    nrooms = 12
    TIMES = []
    day = []
    timeslots = Timeslot.query.filter_by(year=get_current_year(), studyperiod=get_current_studyperiod(),
                                         timetable=get_current_timetable().id).all()
    for timeslot in timeslots:
        TIMES.append(timeslot.day + " " + timeslot.time)
        day.append(timeslot.day)
    day = list(set(day))
    DAYS = {}
    for d in day:
        DAYS[d] = []
    for timeslot in timeslots:
        DAYS[timeslot.day].append(timeslot.day + " " + timeslot.time)
    for d in day:
        DAYS[d] = set(DAYS[d])

    return (STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
            TUTORAVAILABILITY, maxclasssize, minclasssize, nrooms)




def allowed_file(filename):
    '''
    Checks whether the uploaded file has an allowed extension.
    :param filename: The filename to check
    :return: True/False
    '''
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in app.config['ALLOWED_EXTENSIONS']


def upload(file):
    '''
    Save the uploaded file to the UPLOAD_FOLDER directory.

    :param file: The file to upload
    :return: Filename of uploaded file in upload folder
    '''
    if file and allowed_file(file.filename):
        # Make the filename safe, remove unsupported chars
        filename = file.filename
        # Move the file form the temporal folder to
        # the upload folder we setup
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        filename2 = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        return filename2


def checkboxvalue(checkbox):
    '''
    Get value of checkbox.

    :param checkbox: Input from request.form
    :return: 1 if ticked, 0 if not.
    '''
    if (checkbox != None):
        return 1
    else:
        return 0


def read_excel(filename):
    '''
    Read Excel File provided by filename.

    :param filename - path to an Excel file:
    :return: pandas dataframe
    '''
    xl = ExcelFile(filename)
    df = xl.parse(xl.sheet_names[0])
    return df


def get_roll(classid):
    timeclass = TimetabledClass.get(id=classid)
    subject = timeclass.subject
    students = subject.students
    room = timeclass.room
    timeslot = timeclass.timeslot
    return create_roll(students, subject, timeslot, room)



def create_roll(students, subject, timeslot, room):
    document = Document()

    document.add_heading(subject.subname, 0)

    document.add_paragraph('Timeslot: ' + timeslot.day + " " + timeslot.time)
    if room is not None:
        document.add_paragraph('Room: ' + room.name)

    table = document.add_table(rows=1, cols=12)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '4'
    hdr_cells[5].text = '5'
    hdr_cells[6].text = '6'
    hdr_cells[7].text = '7'
    hdr_cells[8].text = '8'
    hdr_cells[9].text = '9'
    hdr_cells[10].text = '10'
    hdr_cells[11].text = '11'
    for item in students:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.name)
    document.save(app.config['UPLOAD_FOLDER'] + '/' + subject.subcode + '.docx')
    return app.config['UPLOAD_FOLDER'] + '/' + subject.subcode + '.docx'
    # document.save('demo.docx')


def create_excel(data):
    writer = pandas.ExcelWriter(app.config['UPLOAD_FOLDER'] + '/timetable.xlsx', engine='xlsxwriter')
    data.to_excel(writer, sheet_name='Timetable')
    writer.save()
    return app.config['UPLOAD_FOLDER'] + '/timetable.xlsx'


def format_timetable_data_for_export():
    timeslots = Timeslot.get_all()
    timeslots = sorted(timeslots, key=attrgetter('daynumeric', 'time'))
    timetable = []
    for i in range(len(timeslots)):
        timeslot = timeslots[i]
        classes = timeslot.timetabledclasses
        for timeclass in classes:
            timetable.append((timeclass.timeslot.day + ' ' + timeclass.timeslot.time, timeclass.subject.subname,
                              timeclass.tutor.name, timeclass.room.name))

    timetable = pandas.DataFrame(timetable)
    timetable.columns = ['Time', 'Subject', 'Tutor', 'Room']
    print(timetable)

    return timetable
